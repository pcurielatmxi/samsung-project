"""
LLM stage implementation using Gemini.

Handles both document extraction (PDF input) and text formatting (JSON input).
Supports multiple document formats: PDF (native upload), DOCX, XLSX (text extraction).
"""

import json
import time
from pathlib import Path
from typing import Any, Optional

from .base import BaseStage, StageResult, FileTask
from ..config import StageConfig
from ..clients.gemini_client import (
    process_document,
    process_document_text,
    process_text_with_document,
    GeminiResponse,
)


# File extensions that need text extraction before processing
TEXT_EXTRACTION_EXTENSIONS = {'.docx', '.doc', '.xlsx', '.xls'}


def extract_docx_text(filepath: Path) -> str:
    """Extract text content from a DOCX file."""
    from docx import Document

    doc = Document(filepath)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n\n".join(parts)


def extract_xlsx_text(filepath: Path) -> str:
    """Extract text content from an XLSX file."""
    import pandas as pd

    parts = []
    xlsx = pd.ExcelFile(filepath)

    for sheet in xlsx.sheet_names:
        df = pd.read_excel(xlsx, sheet_name=sheet)
        if not df.empty:
            parts.append(f"## Sheet: {sheet}\n")
            parts.append(df.to_string(index=False))
            parts.append("")

    return "\n".join(parts)


class LLMStage(BaseStage):
    """
    LLM processing stage.

    - Stage 0: Processes source file (PDF) with document upload
    - Stage N: Processes prior stage's JSON output as text
    """

    def __init__(self, config: StageConfig, config_dir: Path):
        super().__init__(config, config_dir)
        if not config.prompt:
            raise ValueError(f"LLM stage '{config.name}' requires a prompt")

    async def process(
        self,
        task: FileTask,
        input_path: Path,
        enable_enhance: bool = False,
    ) -> StageResult:
        """
        Process a file with Gemini.

        For stage 0 (first stage): Uses document upload for PDF
        For stage N: Reads JSON from prior stage and processes as text

        If enable_enhance and enhance_prompt is configured, runs a second pass
        to review and correct the initial output.
        """
        start_time = time.time()
        total_usage = {"prompt_tokens": 0, "output_tokens": 0, "total_tokens": 0}

        try:
            # Determine if this is first stage (document upload) or later (text)
            if self.config.index == 0:
                # First stage: process source document
                response = await self._process_document(input_path)
            else:
                # Later stage: process prior stage's output as text
                # Pass source_path if include_source is enabled
                source_path = task.source_path if self.config.include_source else None
                response = await self._process_text(input_path, source_path=source_path)

            if not response.success:
                duration_ms = int((time.time() - start_time) * 1000)
                retryable = self._is_retryable_error(response.error)
                return StageResult(
                    success=False,
                    error=response.error,
                    duration_ms=duration_ms,
                    retryable=retryable,
                )

            # Track initial pass usage
            if response.usage:
                total_usage["prompt_tokens"] += response.usage.get("prompt_tokens", 0) or 0
                total_usage["output_tokens"] += response.usage.get("output_tokens", 0) or 0

            result = response.result

            # Enhancement pass: if enabled and enhance_prompt is configured
            if enable_enhance and self.config.has_enhance:
                enhance_response = await self._run_enhancement(result)

                if enhance_response.success:
                    result = enhance_response.result
                    # Add enhancement usage to total
                    if enhance_response.usage:
                        total_usage["prompt_tokens"] += enhance_response.usage.get("prompt_tokens", 0) or 0
                        total_usage["output_tokens"] += enhance_response.usage.get("output_tokens", 0) or 0
                # If enhancement fails, we still return the initial result

            total_usage["total_tokens"] = total_usage["prompt_tokens"] + total_usage["output_tokens"]
            duration_ms = int((time.time() - start_time) * 1000)

            return StageResult(
                success=True,
                result=result,
                usage=total_usage,
                duration_ms=duration_ms,
            )

        except Exception as e:
            duration_ms = int((time.time() - start_time) * 1000)
            return StageResult(
                success=False,
                error=str(e),
                duration_ms=duration_ms,
                retryable=True,
            )

    async def _process_document(self, filepath: Path) -> GeminiResponse:
        """
        Process a document file with Gemini.

        - PDF: Native upload to Gemini
        - DOCX/DOC: Text extraction, then text API
        - XLSX/XLS: Text extraction, then text API
        """
        ext = filepath.suffix.lower()

        if ext in TEXT_EXTRACTION_EXTENSIONS:
            # Extract text from document first
            try:
                if ext in {'.docx', '.doc'}:
                    text = extract_docx_text(filepath)
                elif ext in {'.xlsx', '.xls'}:
                    text = extract_xlsx_text(filepath)
                else:
                    return GeminiResponse(
                        success=False,
                        result=None,
                        error=f"Unsupported format for text extraction: {ext}",
                        model=self.config.model,
                    )

                if not text.strip():
                    return GeminiResponse(
                        success=False,
                        result=None,
                        error=f"No text content extracted from {filepath.name}",
                        model=self.config.model,
                    )

                # Process extracted text
                return process_document_text(
                    text=text,
                    prompt=self.config.prompt,
                    schema=self.config.schema,
                    model=self.config.model,
                )

            except Exception as e:
                return GeminiResponse(
                    success=False,
                    result=None,
                    error=f"Text extraction failed: {e}",
                    model=self.config.model,
                )
        else:
            # PDF or other - use native upload
            return process_document(
                filepath=filepath,
                prompt=self.config.prompt,
                schema=self.config.schema,
                model=self.config.model,
            )

    async def _process_text(
        self,
        input_path: Path,
        source_path: Optional[Path] = None,
    ) -> GeminiResponse:
        """
        Process text content from a prior stage's JSON output.

        Args:
            input_path: Path to prior stage output JSON
            source_path: Optional path to original source document
                        (used when include_source=True)
        """
        # Read prior stage output
        try:
            with open(input_path, "r", encoding="utf-8") as f:
                prior_data = json.load(f)
        except Exception as e:
            return GeminiResponse(
                success=False,
                result=None,
                error=f"Failed to read input file: {e}",
                model=self.config.model,
            )

        # Extract content from prior stage
        content = prior_data.get("content")
        if content is None:
            return GeminiResponse(
                success=False,
                result=None,
                error="Prior stage output has no 'content' field",
                model=self.config.model,
            )

        # If content is a dict (structured), convert to string
        if isinstance(content, dict):
            content = json.dumps(content, indent=2)

        # Use source-aware processing if include_source is enabled
        if self.config.include_source and source_path:
            return process_text_with_document(
                text=content,
                document_path=source_path,
                prompt=self.config.prompt,
                schema=self.config.schema,
                model=self.config.model,
            )

        return process_document_text(
            text=content,
            prompt=self.config.prompt,
            schema=self.config.schema,
            model=self.config.model,
        )

    async def _run_enhancement(self, initial_result: Any) -> GeminiResponse:
        """
        Run enhancement pass on initial extraction result.

        Takes the initial output and runs it through the enhance_prompt
        to review and correct any issues.
        """
        # Prepare the content for enhancement
        if isinstance(initial_result, dict):
            content = json.dumps(initial_result, indent=2)
        else:
            content = str(initial_result)

        # Format the enhance prompt with the initial output
        enhance_prompt = self.config.enhance_prompt.replace("{output_content}", content)

        # Run enhancement through text API
        return process_document_text(
            text=content,
            prompt=enhance_prompt,
            schema=self.config.schema,  # Use same schema for consistent output
            model=self.config.model,
        )

    def _is_retryable_error(self, error: Optional[str]) -> bool:
        """Determine if an error is retryable."""
        if not error:
            return True

        error_lower = error.lower()

        # Non-retryable errors
        non_retryable = [
            "file not found",
            "no api key",
            "invalid api key",
            "exceeds",  # Size/page limit exceeded
            "not supported",
            "invalid format",
        ]

        for term in non_retryable:
            if term in error_lower:
                return False

        return True
