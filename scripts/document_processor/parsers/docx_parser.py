"""Word document parser."""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_docx(file_path: Path) -> str:
    """
    Extract text content from a Word document.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content

    Raises:
        ImportError: If python-docx is not installed
        Exception: If document parsing fails
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document parsing. "
            "Install with: pip install python-docx"
        )

    try:
        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n\n".join(paragraphs)

    except Exception as e:
        logger.error(f"Failed to parse Word document {file_path}: {e}")
        raise
