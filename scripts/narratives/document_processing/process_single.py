#!/usr/bin/env python3
"""
Process a single narrative document (PDF or DOCX) through Gemini extraction.

Handles both PDF (native Gemini upload) and DOCX (text extraction + Gemini text API).
"""

import sys
import json
from pathlib import Path
from datetime import datetime, timezone

# Add project paths
sys.path.insert(0, str(Path(__file__).parent.parent.parent))
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "document_processor_v2"))

from document_processor_v2.gemini_client import process_document, process_document_text, GeminiResponse


def extract_docx_text(filepath: Path) -> str:
    """Extract text content from a DOCX file."""
    from docx import Document

    doc = Document(filepath)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n\n".join(parts)


def extract_xlsx_text(filepath: Path) -> str:
    """Extract text content from an XLSX file."""
    import pandas as pd

    parts = []
    xlsx = pd.ExcelFile(filepath)

    for sheet in xlsx.sheet_names:
        df = pd.read_excel(xlsx, sheet_name=sheet)
        if not df.empty:
            parts.append(f"## Sheet: {sheet}\n")
            parts.append(df.to_string(index=False))
            parts.append("")

    return "\n".join(parts)


def process_narrative(
    filepath: str | Path,
    prompt: str,
    model: str = "gemini-3-flash-preview",
) -> GeminiResponse:
    """
    Process a narrative document (PDF, DOCX, or XLSX).

    Args:
        filepath: Path to document
        prompt: Extraction prompt
        model: Gemini model to use

    Returns:
        GeminiResponse with extraction results
    """
    filepath = Path(filepath)
    suffix = filepath.suffix.lower()

    if suffix == ".pdf":
        # Native PDF processing
        return process_document(filepath, prompt, model=model)

    elif suffix in [".docx", ".doc"]:
        # Extract text from DOCX
        try:
            text = extract_docx_text(filepath)
            if not text.strip():
                return GeminiResponse(
                    success=False,
                    result=None,
                    error="DOCX file is empty or could not extract text",
                    model=model,
                )
            return process_document_text(text, prompt, model=model)
        except Exception as e:
            return GeminiResponse(
                success=False,
                result=None,
                error=f"Failed to extract DOCX: {e}",
                model=model,
            )

    elif suffix in [".xlsx", ".xls"]:
        # Extract text from XLSX
        try:
            text = extract_xlsx_text(filepath)
            if not text.strip():
                return GeminiResponse(
                    success=False,
                    result=None,
                    error="XLSX file is empty or could not extract text",
                    model=model,
                )
            return process_document_text(text, prompt, model=model)
        except Exception as e:
            return GeminiResponse(
                success=False,
                result=None,
                error=f"Failed to extract XLSX: {e}",
                model=model,
            )

    else:
        return GeminiResponse(
            success=False,
            result=None,
            error=f"Unsupported file type: {suffix}",
            model=model,
        )


def save_result(response: GeminiResponse, filepath: Path, output_dir: Path) -> Path:
    """Save extraction result to JSON file."""
    output_dir.mkdir(parents=True, exist_ok=True)

    output_file = output_dir / f"{filepath.stem}.json"

    result = {
        "metadata": {
            "source_file": str(filepath),
            "processed_at": datetime.now(timezone.utc).isoformat(),
            "model": response.model,
            "success": response.success,
            "usage": response.usage,
        },
        "content": response.result if response.success else None,
        "error": response.error,
    }

    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    return output_file


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Process a single narrative document")
    parser.add_argument("filepath", help="Path to document (PDF, DOCX, or XLSX)")
    parser.add_argument("--prompt-file", default="extract_prompt.txt", help="Path to prompt file")
    parser.add_argument("--output-dir", help="Output directory for JSON result")
    parser.add_argument("--model", default="gemini-3-flash-preview", help="Gemini model")

    args = parser.parse_args()

    # Load prompt
    prompt_path = Path(__file__).parent / args.prompt_file
    if not prompt_path.exists():
        print(f"Prompt file not found: {prompt_path}")
        sys.exit(1)

    prompt = prompt_path.read_text(encoding="utf-8").strip()

    # Process document
    filepath = Path(args.filepath)
    print(f"Processing: {filepath.name}")
    print(f"Type: {filepath.suffix}")
    print("-" * 60)

    response = process_narrative(filepath, prompt, model=args.model)

    if response.success:
        print("SUCCESS")
        print(f"Model: {response.model}")
        if response.usage:
            print(f"Usage: {response.usage}")
        print("-" * 60)
        print(response.result)

        # Save if output dir specified
        if args.output_dir:
            output_file = save_result(response, filepath, Path(args.output_dir))
            print(f"\nSaved to: {output_file}")
    else:
        print("FAILED")
        print(f"Error: {response.error}")
        sys.exit(1)
