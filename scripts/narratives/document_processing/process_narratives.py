#!/usr/bin/env python3
"""
Batch process narrative documents (PDF, DOCX, XLSX) through Gemini extraction.

Processes all documents in raw/narratives/ and outputs JSON to processed/narratives/
preserving the subfolder structure.

Usage:
    python process_narratives.py                    # Process all files
    python process_narratives.py --limit 10         # Process first 10 files
    python process_narratives.py --skip-existing    # Skip already processed files
    python process_narratives.py --concurrency 3    # Process 3 files at a time
"""

import sys
import json
import asyncio
import logging
from pathlib import Path
from datetime import datetime, timezone
from dataclasses import dataclass, field
from typing import List, Optional

# Add project paths
# process_narratives.py -> document_processing -> narratives -> scripts -> PROJECT_ROOT
PROJECT_ROOT = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(PROJECT_ROOT / "scripts"))

from src.config.settings import Settings
from document_processor_v2.gemini_client import process_document, process_document_text, GeminiResponse

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

# Paths
RAW_NARRATIVES_DIR = Settings.RAW_DATA_DIR / "narratives"
PROCESSED_NARRATIVES_DIR = Settings.PROCESSED_DATA_DIR / "narratives"
PROMPT_FILE = Path(__file__).parent / "extract_prompt.txt"

# Supported extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls"}

# Default model
DEFAULT_MODEL = "gemini-3-flash-preview"


@dataclass
class ProcessingStats:
    """Track processing statistics."""
    total_files: int = 0
    processed: int = 0
    skipped_existing: int = 0
    skipped_unsupported: int = 0
    errors: int = 0
    successful: int = 0
    total_tokens: int = 0
    error_files: List[tuple] = field(default_factory=list)


def extract_docx_text(filepath: Path) -> str:
    """Extract text content from a DOCX file."""
    from docx import Document

    doc = Document(filepath)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n\n".join(parts)


def extract_xlsx_text(filepath: Path) -> str:
    """Extract text content from an XLSX file."""
    import pandas as pd

    parts = []
    xlsx = pd.ExcelFile(filepath)

    for sheet in xlsx.sheet_names:
        df = pd.read_excel(xlsx, sheet_name=sheet)
        if not df.empty:
            parts.append(f"## Sheet: {sheet}\n")
            parts.append(df.to_string(index=False))
            parts.append("")

    return "\n".join(parts)


def process_narrative_sync(
    filepath: Path,
    prompt: str,
    model: str = DEFAULT_MODEL,
) -> GeminiResponse:
    """
    Process a narrative document (PDF, DOCX, or XLSX).

    Args:
        filepath: Path to document
        prompt: Extraction prompt
        model: Gemini model to use

    Returns:
        GeminiResponse with extraction results
    """
    suffix = filepath.suffix.lower()

    if suffix == ".pdf":
        # Native PDF processing
        return process_document(filepath, prompt, model=model)

    elif suffix in [".docx", ".doc"]:
        # Extract text from DOCX
        try:
            text = extract_docx_text(filepath)
            if not text.strip():
                return GeminiResponse(
                    success=False,
                    result=None,
                    error="DOCX file is empty or could not extract text",
                    model=model,
                )
            return process_document_text(text, prompt, model=model)
        except Exception as e:
            return GeminiResponse(
                success=False,
                result=None,
                error=f"Failed to extract DOCX: {e}",
                model=model,
            )

    elif suffix in [".xlsx", ".xls"]:
        # Extract text from XLSX
        try:
            text = extract_xlsx_text(filepath)
            if not text.strip():
                return GeminiResponse(
                    success=False,
                    result=None,
                    error="XLSX file is empty or could not extract text",
                    model=model,
                )
            return process_document_text(text, prompt, model=model)
        except Exception as e:
            return GeminiResponse(
                success=False,
                result=None,
                error=f"Failed to extract XLSX: {e}",
                model=model,
            )

    else:
        return GeminiResponse(
            success=False,
            result=None,
            error=f"Unsupported file type: {suffix}",
            model=model,
        )


def save_result(
    response: GeminiResponse,
    source_path: Path,
    output_path: Path,
    error_path: Path,
) -> None:
    """Save extraction result to .extract.json or .extract.error.json file."""
    output_path.parent.mkdir(parents=True, exist_ok=True)

    result = {
        "metadata": {
            "source_file": str(source_path),
            "processed_at": datetime.now(timezone.utc).isoformat(),
            "model": response.model,
            "stage": "extract",
            "success": response.success,
            "usage": response.usage,
        },
        "content": response.result if response.success else None,
        "error": response.error,
    }

    # Save to appropriate file based on success/failure
    if response.success:
        # Remove error file if exists (retry succeeded)
        if error_path.exists():
            error_path.unlink()
        target_path = output_path
    else:
        target_path = error_path

    with open(target_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)


def discover_files(input_dir: Path) -> List[Path]:
    """Discover all processable files recursively."""
    files = []
    for ext in SUPPORTED_EXTENSIONS:
        files.extend(input_dir.rglob(f"*{ext}"))
        # Also check uppercase extensions
        files.extend(input_dir.rglob(f"*{ext.upper()}"))
    return sorted(set(files))


def get_output_paths(source_path: Path, input_dir: Path, output_dir: Path) -> tuple[Path, Path]:
    """Get output paths (extract and error) preserving subfolder structure.

    Uses same naming convention as document_processor_v2/pipeline.py:
    - {stem}.extract.json - Stage 1 extraction output
    - {stem}.extract.error.json - Stage 1 error marker
    """
    rel_path = source_path.relative_to(input_dir)
    # Use string concatenation instead of with_suffix to handle filenames with periods
    base_dir = output_dir / rel_path.parent
    stem = source_path.stem
    return base_dir / f"{stem}.extract.json", base_dir / f"{stem}.extract.error.json"


async def process_file(
    source_path: Path,
    output_path: Path,
    error_path: Path,
    prompt: str,
    model: str,
    semaphore: asyncio.Semaphore,
    stats: ProcessingStats,
) -> None:
    """Process a single file with rate limiting."""
    async with semaphore:
        try:
            # Run sync processing in thread pool
            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None,
                process_narrative_sync,
                source_path,
                prompt,
                model,
            )

            save_result(response, source_path, output_path, error_path)

            if response.success:
                stats.successful += 1
                if response.usage:
                    stats.total_tokens += response.usage.get("total_tokens", 0)
                logger.info(f"✓ {source_path.name}")
            else:
                stats.errors += 1
                stats.error_files.append((source_path, response.error))
                logger.error(f"✗ {source_path.name}: {response.error}")

        except Exception as e:
            stats.errors += 1
            stats.error_files.append((source_path, str(e)))
            logger.error(f"✗ {source_path.name}: {e}")

        stats.processed += 1


async def process_batch(
    files: List[tuple],
    prompt: str,
    model: str,
    concurrency: int,
    stats: ProcessingStats,
) -> None:
    """Process a batch of files with concurrency control."""
    semaphore = asyncio.Semaphore(concurrency)

    tasks = [
        process_file(source, output, error, prompt, model, semaphore, stats)
        for source, output, error in files
    ]

    await asyncio.gather(*tasks)


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Batch process narrative documents")
    parser.add_argument("--limit", type=int, help="Limit number of files to process")
    parser.add_argument("--force", action="store_true", help="Reprocess all files (ignore existing)")
    parser.add_argument("--retry-errors", action="store_true", help="Retry only files that previously failed")
    parser.add_argument("--concurrency", type=int, default=3, help="Number of concurrent requests")
    parser.add_argument("--model", default=DEFAULT_MODEL, help="Gemini model to use")
    parser.add_argument("--dry-run", action="store_true", help="Show what would be processed")

    args = parser.parse_args()

    # Load prompt
    if not PROMPT_FILE.exists():
        logger.error(f"Prompt file not found: {PROMPT_FILE}")
        sys.exit(1)

    prompt = PROMPT_FILE.read_text(encoding="utf-8").strip()

    # Check input directory
    if not RAW_NARRATIVES_DIR.exists():
        logger.error(f"Input directory not found: {RAW_NARRATIVES_DIR}")
        logger.info("Please create the directory and add narrative documents.")
        sys.exit(1)

    # Discover files
    all_files = discover_files(RAW_NARRATIVES_DIR)
    logger.info(f"Found {len(all_files)} files in {RAW_NARRATIVES_DIR}")

    # Build file list with output paths and apply idempotency
    stats = ProcessingStats(total_files=len(all_files))
    files_to_process = []
    skipped_completed = 0
    skipped_failed = 0

    for source_path in all_files:
        output_path, error_path = get_output_paths(source_path, RAW_NARRATIVES_DIR, PROCESSED_NARRATIVES_DIR)

        # Idempotency logic (default: skip completed files)
        if not args.force:
            if output_path.exists():
                # Already successfully processed
                if args.retry_errors:
                    # Only retrying errors, skip completed
                    skipped_completed += 1
                    continue
                else:
                    # Normal run, skip completed
                    skipped_completed += 1
                    continue
            elif error_path.exists():
                # Previously failed
                if not args.retry_errors:
                    # Normal run, skip errors too (already tried)
                    skipped_failed += 1
                    continue
                # retry-errors mode: will reprocess this file

        files_to_process.append((source_path, output_path, error_path))

    # Apply limit
    if args.limit:
        files_to_process = files_to_process[:args.limit]

    # Log status
    logger.info(f"Files to process: {len(files_to_process)}")
    if skipped_completed:
        logger.info(f"Skipped (completed): {skipped_completed}")
    if skipped_failed:
        logger.info(f"Skipped (failed): {skipped_failed}")

    if args.dry_run:
        logger.info("\nDry run - files that would be processed:")
        for source, output, error in files_to_process[:20]:
            rel_path = source.relative_to(RAW_NARRATIVES_DIR)
            status = "[retry]" if error.exists() else "[new]"
            logger.info(f"  {status} {rel_path}")
        if len(files_to_process) > 20:
            logger.info(f"  ... and {len(files_to_process) - 20} more")
        return

    if not files_to_process:
        logger.info("No files to process.")
        return

    # Process files
    logger.info(f"\nProcessing with concurrency={args.concurrency}, model={args.model}")
    logger.info("-" * 60)

    asyncio.run(process_batch(
        files_to_process,
        prompt,
        args.model,
        args.concurrency,
        stats,
    ))

    # Print summary
    logger.info("-" * 60)
    logger.info("Processing complete!")
    logger.info(f"  Successful: {stats.successful}")
    logger.info(f"  Errors: {stats.errors}")
    logger.info(f"  Total tokens: {stats.total_tokens:,}")

    if stats.error_files:
        logger.info("\nFailed files:")
        for path, error in stats.error_files[:10]:
            logger.info(f"  {path.name}: {error[:80]}")
        if len(stats.error_files) > 10:
            logger.info(f"  ... and {len(stats.error_files) - 10} more errors")


if __name__ == "__main__":
    main()
