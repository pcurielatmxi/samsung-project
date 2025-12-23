#!/usr/bin/env python3
"""
TBM Analysis DOCX Generator
===========================

Generates styled DOCX reports with MXI branding from processed report content.
Input: JSON file from tbm_data_processor.py (with optional narratives)
Output: Professionally styled Word document

Usage:
    python docx_generator.py --input report_content.json --output report.docx
"""

import argparse
import io
import json
import requests
from datetime import datetime
from pathlib import Path
from typing import Optional

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# MXI BRAND COLORS
# =============================================================================

class Colors:
    PRIMARY_BLUE = RGBColor(30, 58, 95)      # Headers, titles
    SECONDARY_BLUE = RGBColor(46, 89, 132)   # Subheaders
    ACCENT_BLUE = RGBColor(74, 144, 217)     # Category headers
    CRITICAL_RED = RGBColor(196, 30, 58)     # Alerts, idle
    WARNING_ORANGE = RGBColor(230, 126, 34)  # Warnings
    SUCCESS_GREEN = RGBColor(39, 174, 96)    # Good metrics
    DARK_GRAY = RGBColor(44, 62, 80)         # Body text
    WHITE = RGBColor(255, 255, 255)


# =============================================================================
# STYLING HELPERS
# =============================================================================

def set_spacing(paragraph, before=0, after=0, line=1.0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_cell_shading(cell, hex_color: str):
    """Add background color to table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    cell._element.get_or_add_tcPr().append(shading)


def add_cell_padding(cell, padding: int = 100):
    """Add padding to table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for side in ['top', 'bottom', 'left', 'right']:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(padding))
        margin.set(qn('w:type'), 'dxa')
        tc_mar.append(margin)
    tc_pr.append(tc_mar)


def set_table_borders(table, color: str = '000000', size: int = 4):
    """Add borders to all cells in table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        border.set(qn('w:space'), '0')
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def download_image(url: str, timeout: int = 10) -> Optional[io.BytesIO]:
    """Download image from URL and return as BytesIO buffer."""
    try:
        response = requests.get(url, timeout=timeout)
        if response.status_code == 200:
            buf = io.BytesIO(response.content)
            buf.seek(0)
            return buf
    except Exception as e:
        print(f"  Warning: Could not download image: {e}")
    return None


# =============================================================================
# DOCUMENT BUILDER
# =============================================================================

class TBMReportBuilder:
    """Builds the TBM Analysis Report DOCX."""

    def __init__(self, content: dict):
        self.content = content
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """Configure document margins and default paragraph style."""
        for section in self.doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Set default paragraph style: 0 before, 0 after, single line spacing
        style = self.doc.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

    def build(self, output_path: str):
        """Build complete report."""
        self._add_cover_page()
        self.doc.add_page_break()

        self._add_executive_dashboard()
        self.doc.add_page_break()

        self._add_category_1()
        self._add_category_2()
        self._add_category_3()
        self._add_category_4()
        self._add_category_5()
        self._add_category_6()
        self._add_category_7()
        self._add_category_8()

        self._add_end_marker()

        self.doc.save(output_path)

    # =========================================================================
    # COVER PAGE
    # =========================================================================

    def _add_cover_page(self):
        """Build cover page."""
        # Large top spacer
        p = self.doc.add_paragraph()
        set_spacing(p, before=120)

        # MXI
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('MXI')
        run.font.size = Pt(48)
        run.font.bold = True
        run.font.color.rgb = Colors.PRIMARY_BLUE

        # Tagline
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Document Control • Project Management • Field Services')
        run.font.size = Pt(11)
        run.font.color.rgb = Colors.DARK_GRAY

        # Spacer before title
        p = self.doc.add_paragraph()
        set_spacing(p, before=36)

        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('TBM ANALYSIS REPORT')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = Colors.PRIMARY_BLUE

        # Subtitle
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{self.content['contractor_group']} Subcontractor Submittals")
        run.font.size = Pt(14)

        # Spacer before project info
        p = self.doc.add_paragraph()
        set_spacing(p, before=24)

        # Project info
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('SAMSUNG TAYLOR FAB - SECAI')
        run.font.size = Pt(13)
        run.font.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('1530 FM973, Taylor, TX 76574')
        run.font.size = Pt(11)

        # Spacer before table
        p = self.doc.add_paragraph()
        set_spacing(p, before=24)

        # Summary stats table
        table = self.doc.add_table(rows=2, cols=4)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_table_borders(table)

        headers = ['Total Tasks', 'Contractors', 'Floors', 'Report Pages']
        values = [
            str(self.content['total_tasks']),
            str(self.content['total_contractors']),
            str(self.content['total_floors']),
            '—'
        ]

        for i, (h, v) in enumerate(zip(headers, values)):
            # Header row - gray background
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            add_cell_shading(cell, 'D9D9D9')

            # Value row - no fill, red text
            cell = table.rows[1].cells[i]
            cell.text = v
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(14)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = Colors.CRITICAL_RED

        # Spacer before date
        p = self.doc.add_paragraph()
        set_spacing(p, before=24)

        # Date
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.content['report_date_formatted'])
        run.font.size = Pt(13)
        run.font.bold = True

        # Large spacer before footer
        for _ in range(5):
            self.doc.add_paragraph()

        # Boots to Bytes
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Boots to Bytes')
        run.font.size = Pt(11)
        run.font.italic = True

        # Spacer
        p = self.doc.add_paragraph()
        set_spacing(p, before=24)

        # Confidentiality
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('CONFIDENTIAL - FOR INTERNAL USE ONLY')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = Colors.CRITICAL_RED

    # =========================================================================
    # EXECUTIVE DASHBOARD
    # =========================================================================

    def _add_executive_dashboard(self):
        """Build executive dashboard."""
        p = self.doc.add_paragraph()
        run = p.add_run('Executive Dashboard')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = Colors.PRIMARY_BLUE

        # Spacer after title
        p = self.doc.add_paragraph()
        set_spacing(p, before=12)

        # KPI cards (2 rows × 3 cols)
        self._add_kpi_cards()

        # Spacer before TBM Summary
        p = self.doc.add_paragraph()
        set_spacing(p, before=18)

        # TBM Summary
        self._add_tbm_summary_section()

        # Spacer before chart
        p = self.doc.add_paragraph()
        set_spacing(p, before=18)

        # Bar chart - Planned vs Actual by location
        self._add_chart()

    def _get_kpi_font_color(self, kpi_type: str, value: float) -> RGBColor:
        """Get font color for KPI based on thresholds."""
        if kpi_type == 'lpi':
            # LPI: <40% red, 40-70% orange, >70% green
            if value < 40:
                return Colors.CRITICAL_RED
            elif value < 70:
                return Colors.WARNING_ORANGE
            else:
                return Colors.SUCCESS_GREEN
        elif kpi_type == 'dilr':
            # DILR: >35% red, 25-35% orange, <25% green (lower is better)
            if value > 35:
                return Colors.CRITICAL_RED
            elif value > 25:
                return Colors.WARNING_ORANGE
            else:
                return Colors.SUCCESS_GREEN
        elif kpi_type == 'idle':
            # Idle hours: always show as concerning (red)
            return Colors.CRITICAL_RED
        elif kpi_type == 'compliance':
            # Same as LPI thresholds
            if value < 40:
                return Colors.CRITICAL_RED
            elif value < 70:
                return Colors.WARNING_ORANGE
            else:
                return Colors.SUCCESS_GREEN
        elif kpi_type == 'zero_verification':
            # High zero verification is bad: >50% red, 30-50% orange, <30% green
            if value > 50:
                return Colors.CRITICAL_RED
            elif value > 30:
                return Colors.WARNING_ORANGE
            else:
                return Colors.SUCCESS_GREEN
        elif kpi_type == 'projection':
            # Monthly projection - neutral dark gray
            return Colors.DARK_GRAY
        return Colors.DARK_GRAY

    def _add_kpi_cards(self):
        """Build KPI cards table - 2 rows of 3 cards each."""
        # Row 1 KPIs
        table1 = self.doc.add_table(rows=3, cols=3)
        set_table_borders(table1)

        lpi = self.content['lpi']
        dilr = self.content['dilr']
        idle_hours = self.content['idle_time_hours']

        row1_cards = [
            ('Labor Productivity Index (LPI)', f"{lpi:.0f}%", f"Target: {self.content['lpi_target']}%", 'lpi', lpi),
            ('Direct-Indirect Labor Ratio (DILR)', f"{dilr:.0f}%", f"Target: <{self.content['dilr_target']}%", 'dilr', dilr),
            ('Idle Time Hours', f"{idle_hours:.0f}", f"@ ${self.content['labor_rate']}/hr = ${self.content['idle_time_cost']:,.0f}", 'idle', idle_hours),
        ]

        for i, (title, value, note, kpi_type, kpi_val) in enumerate(row1_cards):
            font_color = self._get_kpi_font_color(kpi_type, kpi_val)

            # Header row - gray background
            cell = table1.rows[0].cells[i]
            cell.text = title
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.bold = True
            add_cell_shading(cell, 'D9D9D9')

            # Value row - colored font (no background)
            cell = table1.rows[1].cells[i]
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = font_color

            # Note row
            cell = table1.rows[2].cells[i]
            cell.text = note
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(11)

        # Row 2 KPIs
        self.doc.add_paragraph()
        table2 = self.doc.add_table(rows=3, cols=3)
        set_table_borders(table2)

        compliance = self.content.get('tbm_compliance_rate', self.content['lpi'])
        zero_rate = self.content['zero_verification_rate']
        monthly_proj = self.content['monthly_projection']

        row2_cards = [
            ('TBM Compliance Rate', f"{compliance:.0f}%", f"TBM Received {self.content.get('tbm_received_time', 'Late')}", 'compliance', compliance),
            ('Zero Verification Rate', f"{zero_rate:.0f}%", f"{self.content['zero_verification_count']} of {self.content['total_locations']} locations", 'zero_verification', zero_rate),
            ('Monthly Projection', f"${monthly_proj:.0f}K", '26 working days', 'projection', monthly_proj),
        ]

        for i, (title, value, note, kpi_type, kpi_val) in enumerate(row2_cards):
            font_color = self._get_kpi_font_color(kpi_type, kpi_val)

            # Header row - gray background
            cell = table2.rows[0].cells[i]
            cell.text = title
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.bold = True
            add_cell_shading(cell, 'D9D9D9')

            # Value row - colored font (no background)
            cell = table2.rows[1].cells[i]
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = font_color

            # Note row
            cell = table2.rows[2].cells[i]
            cell.text = note
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    def _add_tbm_summary_section(self):
        """Build TBM Summary table."""
        p = self.doc.add_paragraph()
        run = p.add_run('TBM Summary')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = Colors.SECONDARY_BLUE

        metrics = self.content['contractor_metrics']
        table = self.doc.add_table(rows=len(metrics) + 2, cols=5)
        set_table_borders(table)

        # Headers - gray background
        headers = ['Contractor', 'Planned MP', 'TBM Presence', 'Actual', 'Accuracy']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            add_cell_shading(cell, 'D9D9D9')

        # Data rows
        for idx, m in enumerate(metrics, 1):
            # Contractor name - bold
            cell = table.rows[idx].cells[0]
            cell.text = m['name']
            cell.paragraphs[0].runs[0].font.bold = True

            table.rows[idx].cells[1].text = f"{m['total_planned']:.0f}"
            table.rows[idx].cells[2].text = f"{m['tbm_presence']:.0f}"
            table.rows[idx].cells[3].text = f"{m['total_verified']:.0f}"
            table.rows[idx].cells[4].text = f"{m['accuracy']:.0f}%"

            # Color code accuracy with font color
            acc_cell = table.rows[idx].cells[4]
            acc = m['accuracy']
            acc_cell.paragraphs[0].runs[0].font.bold = True
            if acc < 40:
                acc_cell.paragraphs[0].runs[0].font.color.rgb = Colors.CRITICAL_RED
            elif acc < 70:
                acc_cell.paragraphs[0].runs[0].font.color.rgb = Colors.WARNING_ORANGE
            else:
                acc_cell.paragraphs[0].runs[0].font.color.rgb = Colors.SUCCESS_GREEN

        # Total row
        total_idx = len(metrics) + 1
        table.rows[total_idx].cells[0].text = 'TOTAL'
        table.rows[total_idx].cells[1].text = f"{self.content['total_planned']:.0f}"
        table.rows[total_idx].cells[2].text = f"{self.content['total_tbm_presence']:.0f}"
        table.rows[total_idx].cells[3].text = f"{self.content['total_verified']:.0f}"
        table.rows[total_idx].cells[4].text = f"{self.content['total_accuracy']:.0f}%"

        for cell in table.rows[total_idx].cells:
            cell.paragraphs[0].runs[0].font.bold = True

    def _add_chart(self):
        """Generate and embed Planned vs Actual bar chart using matplotlib."""
        # Collect data for chart
        tasks = self.content.get('task_details', [])

        if not tasks:
            # Fallback: skip chart if no task details
            return

        # Group by contractor and location
        chart_data = []
        for task in tasks:
            chart_data.append({
                'contractor': task.get('company', 'Unknown'),
                'location': task.get('location', 'Unknown'),
                'level': task.get('level', ''),
                'planned': task.get('planned', 0),
                'actual': task.get('actual', 0),
            })

        # Sort by contractor then location
        chart_data.sort(key=lambda x: (x['contractor'], x['location']))

        # Limit to first 25 locations for readability
        chart_data = chart_data[:25]

        if not chart_data:
            return

        # Create figure with higher resolution
        fig, ax = plt.subplots(figsize=(10, 5))

        locations = [f"{d['location']}" for d in chart_data]
        planned = [d['planned'] for d in chart_data]
        actual = [d['actual'] for d in chart_data]

        x = range(len(locations))
        width = 0.35

        # MXI colors
        bars1 = ax.bar([i - width/2 for i in x], planned, width, label='Planned per TBM', color='#1E3A5F')
        bars2 = ax.bar([i + width/2 for i in x], actual, width, label='Actual per TBM', color='#C41E3A')

        ax.set_ylabel('Workers', fontsize=11)
        ax.set_xticks(x)
        ax.set_xticklabels(locations, rotation=45, ha='right', fontsize=9)
        ax.legend(fontsize=10)
        ax.tick_params(axis='y', labelsize=10)

        # Add contractor labels at bottom and find group boundaries
        contractors = [d['contractor'] for d in chart_data]
        prev_contractor = None
        contractor_starts = []
        for i, c in enumerate(contractors):
            if c != prev_contractor:
                contractor_starts.append((i, c))
                prev_contractor = c

        # Add vertical separator lines between contractor groups
        for idx, (start_idx, contractor) in enumerate(contractor_starts):
            if idx > 0:
                # Add vertical line at boundary between contractors
                # Line position is 0.5 before the start of the new contractor
                ax.axvline(x=start_idx - 0.5, color='#888888', linestyle='--', linewidth=1.5, alpha=0.7)

        # Add contractor group labels
        for idx, (start_idx, contractor) in enumerate(contractor_starts):
            # Find end of this contractor's data
            end_idx = start_idx
            for j in range(start_idx, len(contractors)):
                if contractors[j] == contractor:
                    end_idx = j
                else:
                    break
            mid = (start_idx + end_idx) / 2
            ax.text(mid, -0.18, contractor, ha='center', va='top',
                    transform=ax.get_xaxis_transform(), fontsize=11, fontweight='bold')

        plt.tight_layout()

        # Save to bytes buffer with higher DPI for better resolution
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        buf.seek(0)
        plt.close(fig)

        # Add to document
        self.doc.add_picture(buf, width=Inches(6.5))
        buf.close()

    # =========================================================================
    # CATEGORIES
    # =========================================================================

    def _add_category_header(self, number: int, title: str):
        """Add category header."""
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        run = p.add_run(f'Category {number}: {title}')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = Colors.ACCENT_BLUE

    def _add_subheader(self, text: str):
        """Add subheader."""
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = Colors.SECONDARY_BLUE

    def _add_alert_box(self, text: str, bg_hex: str = 'FADBD8', text_color=None):
        """Add alert box."""
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        add_cell_shading(cell, bg_hex)
        add_cell_padding(cell, 150)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = text_color or Colors.CRITICAL_RED

    def _add_observation_box(self, text: str):
        """Add shaded observation box."""
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        add_cell_shading(cell, 'ECF0F1')
        add_cell_padding(cell, 100)

        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = Colors.DARK_GRAY

    def _add_category_1(self):
        """Category 1: TBM Planning & Documentation Performance."""
        self._add_category_header(1, 'TBM Planning & Documentation Performance')

        narratives = self.content.get('narratives', {}).get('cat1', {})

        # Opening narrative
        better_contractor = max(self.content['contractor_metrics'], key=lambda x: x['accuracy'])
        opening = narratives.get('opening',
            f"TBM verification shows improved performance with {self.content['lpi']:.0f}% overall accuracy. "
            f"{better_contractor['name']} continues to outperform in workforce documentation.")

        p = self.doc.add_paragraph()
        p.add_run(opening)

        self.doc.add_paragraph()

        # Calculate unaccounted workers
        total_planned = self.content['total_planned']
        total_verified = self.content['total_verified']
        unaccounted = total_planned - total_verified
        unaccounted_pct = (unaccounted / total_planned * 100) if total_planned > 0 else 0

        # Unaccounted workers alert
        self._add_alert_box(
            f"{unaccounted:.0f} OF {total_planned:.0f} WORKERS ({unaccounted_pct:.0f}%) UNACCOUNTED - Documentation gap persists but improving",
            'FADBD8', Colors.CRITICAL_RED
        )

        self.doc.add_paragraph()

        # TBM Timing Issue alert (if timing data available)
        tbm_received = self.content.get('tbm_received_time', '1:22 PM')
        workers_departing = self.content.get('workers_departing_time', '2:30 PM')
        self._add_alert_box(
            f"TBM TIMING ISSUE: TBM received at {tbm_received} with workers leaving at {workers_departing} - insufficient time for proper analysis",
            'FDEBD0', Colors.WARNING_ORANGE
        )

        # Subheader 1.1
        self._add_subheader('1.1 TBM Submission Timing Impact')

        # Timing narrative
        timing_narrative = self.content.get('narratives', {}).get('cat1', {}).get('timing_narrative',
            f"The Yates TBM was received by MXI at {tbm_received} on {self.content['report_date_formatted']}. "
            f"Some workers began cutting short and leaving the site by {workers_departing}, providing MXI field verification team with only approximately 1 hour to "
            f"review the TBM, plan verification routes, and conduct field observations.")
        p = self.doc.add_paragraph()
        p.add_run(timing_narrative)

        self.doc.add_paragraph()

        # Timing constraint explanation
        total_locations = self.content['total_locations']
        total_workers = self.content['total_planned']
        verification_window = self.content.get('verification_window_minutes', 68)
        minutes_per_location = verification_window / total_locations if total_locations > 0 else 0

        constraint_text = (
            f"This timing constraint significantly impacts verification accuracy. MXI cannot effectively verify {total_workers:.0f} "
            f"committed workers across {total_locations} locations in {verification_window} minutes. Late TBM submission combined with early worker "
            f"departure creates a systemic verification gap that undermines the entire TBM process.")
        self._add_observation_box(constraint_text)

        self.doc.add_paragraph()

        # Event timing table
        table = self.doc.add_table(rows=6, cols=3)
        set_table_borders(table)

        headers = ['Event', 'Time', 'Impact']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            add_cell_shading(cell, 'D9D9D9')

        timing_rows = [
            ('TBM Received', tbm_received, 'Late submission'),
            ('Workers Departing', workers_departing, f'{verification_window} minutes available'),
            ('Total Locations', str(total_locations), f'{minutes_per_location:.1f} min per location'),
            ('Total Workers', f'{total_workers:.0f}', f'{verification_window * 60 / total_workers:.0f} seconds per worker' if total_workers > 0 else 'N/A'),
            ('Verification Window', f'{verification_window} minutes', 'INSUFFICIENT'),
        ]

        for idx, (event, time, impact) in enumerate(timing_rows, 1):
            table.rows[idx].cells[0].text = event
            table.rows[idx].cells[1].text = time
            table.rows[idx].cells[2].text = impact
            # Highlight insufficient
            if 'INSUFFICIENT' in impact:
                table.rows[idx].cells[2].paragraphs[0].runs[0].font.bold = True
                table.rows[idx].cells[2].paragraphs[0].runs[0].font.color.rgb = Colors.CRITICAL_RED

        self.doc.add_paragraph()

        # Operational Impact paragraph
        p = self.doc.add_paragraph()
        run = p.add_run('OPERATIONAL IMPACT: ')
        run.font.bold = True
        p.add_run(
            f"The {verification_window}-minute verification window makes comprehensive verification "
            f"mathematically impossible. With {total_locations} locations to verify, MXI has an average of {minutes_per_location:.1f} minutes per location - "
            f"insufficient time to walk to location, observe workers, document findings, and move to next area.")

        self.doc.add_paragraph()

        # Contractor accountability table
        metrics = self.content['contractor_metrics']
        table = self.doc.add_table(rows=len(metrics) + 2, cols=5)
        set_table_borders(table)

        headers = ['Contractor', 'Planned', 'Verified', 'Unaccounted', '% Unaccounted']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            add_cell_shading(cell, 'D9D9D9')

        for idx, m in enumerate(metrics, 1):
            unacct = m['total_planned'] - m['total_verified']
            unacct_pct = (unacct / m['total_planned'] * 100) if m['total_planned'] > 0 else 0

            table.rows[idx].cells[0].text = m['name']
            table.rows[idx].cells[1].text = f"{m['total_planned']:.0f}"
            table.rows[idx].cells[2].text = f"{m['total_verified']:.0f}"
            table.rows[idx].cells[3].text = f"{unacct:.0f}"
            table.rows[idx].cells[4].text = f"{unacct_pct:.0f}%"

        # Total row
        total_idx = len(metrics) + 1
        table.rows[total_idx].cells[0].text = 'TOTAL'
        table.rows[total_idx].cells[1].text = f"{total_planned:.0f}"
        table.rows[total_idx].cells[2].text = f"{total_verified:.0f}"
        table.rows[total_idx].cells[3].text = f"{unaccounted:.0f}"
        table.rows[total_idx].cells[4].text = f"{unaccounted_pct:.0f}%"
        for cell in table.rows[total_idx].cells:
            cell.paragraphs[0].runs[0].font.bold = True

    def _add_category_2(self):
        """Category 2: Zero Verification Locations."""
        self._add_category_header(2, 'Zero Verification Locations')

        narratives = self.content.get('narratives', {}).get('cat2', {})
        intro = narratives.get('intro',
            'The following locations had TBM commitments but zero field verification. This represents planning/documentation gaps.')

        p = self.doc.add_paragraph()
        p.add_run(intro)

        self.doc.add_paragraph()

        # Calculate total unverified workers from zero verification locations
        zero_locs = self.content['zero_verification_locations']
        total_zero_workers = sum(z['tbm_planned'] for z in zero_locs)
        count = self.content['zero_verification_count']

        # Alert box with worker count
        self._add_alert_box(
            f"{count} LOCATIONS WITH 0% VERIFICATION - {total_zero_workers:.0f} workers unaccounted (documentation gap, NO COST)",
            'FADBD8', Colors.CRITICAL_RED
        )

        # Group by contractor with numbered subheaders
        contractors = []
        for z in zero_locs:
            if z['company'] not in contractors:
                contractors.append(z['company'])

        for sub_num, contractor in enumerate(contractors, 1):
            contractor_locs = [z for z in zero_locs if z['company'] == contractor]
            total_workers = sum(z['tbm_planned'] for z in contractor_locs)

            self._add_subheader(f"2.{sub_num} {contractor} - Zero Verification ({len(contractor_locs)} locations, {total_workers:.0f} workers)")

            # Table with borders
            table = self.doc.add_table(rows=len(contractor_locs) + 1, cols=4)
            set_table_borders(table)

            headers = ['Floor', 'Location', 'TBM', 'Verified']
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                add_cell_shading(cell, 'D9D9D9')

            for idx, loc in enumerate(contractor_locs, 1):
                table.rows[idx].cells[0].text = loc['floor'] if loc['floor'] else '-'
                table.rows[idx].cells[1].text = loc['location']
                table.rows[idx].cells[2].text = f"{loc['tbm_planned']:.0f}"
                table.rows[idx].cells[3].text = '0'

    def _add_category_3(self):
        """Category 3: Idle Worker Observations - ACTUAL COST."""
        self._add_category_header(3, 'Idle Worker Observations - ACTUAL COST')

        narratives = self.content.get('narratives', {}).get('cat3', {})

        # Methodology
        methodology = narratives.get('methodology', f"CALCULATION METHODOLOGY: Idle time cost based on 1-hour idle duration per observation. Labor rate: ${self.content['labor_rate']}/hr.")
        self._add_alert_box(methodology, 'FDEBD0', Colors.WARNING_ORANGE)

        self.doc.add_paragraph()

        intro = narratives.get('intro', 'Field verification identified workers physically observed idle. This represents ACTUAL WASTE.')
        p = self.doc.add_paragraph()
        p.add_run(intro)

        self.doc.add_paragraph()

        # Alert with totals
        total_obs = len(self.content['idle_observations'])
        total_hours = self.content['idle_time_hours']
        total_cost = self.content['idle_time_cost']
        self._add_alert_box(f"IDLE TIME: {total_obs} OBSERVATIONS - {total_hours:.0f} WORKER-HOURS - ${total_cost:,.0f} DAILY WASTE")

        # Group by contractor
        idle_obs = self.content['idle_observations']
        contractors = set(o['company'] for o in idle_obs)

        for contractor in contractors:
            contractor_obs = [o for o in idle_obs if o['company'] == contractor]
            self._add_subheader(f"{contractor} - Idle Time Observations ({len(contractor_obs)} locations)")

            for obs in contractor_obs:
                self.doc.add_paragraph()

                # Task header
                p = self.doc.add_paragraph()
                header = f"{obs['task_id']}: {obs['location']}"
                if obs['level']:
                    header += f" ({obs['level']})"
                header += f" - {obs['status']}"
                run = p.add_run(header)
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = Colors.CRITICAL_RED

                # Observation
                if obs['observation_text']:
                    self._add_observation_box(f"Field Observation ({obs['inspector']}): {obs['observation_text']}")

                # Analysis
                p = self.doc.add_paragraph()
                run = p.add_run('Analysis: ')
                run.font.bold = True
                p.add_run(obs['analysis'])

                # Photos - embed first photo if available
                photo_urls = obs.get('photo_urls', [])
                if photo_urls:
                    self.doc.add_paragraph()
                    # Download and embed first photo (limit to 1 per observation)
                    for url in photo_urls[:1]:
                        print(f"  Downloading photo for task {obs['task_id']}...")
                        img_buf = download_image(url)
                        if img_buf:
                            try:
                                self.doc.add_picture(img_buf, width=Inches(4.0))
                                img_buf.close()
                            except Exception as e:
                                print(f"    Warning: Could not embed image: {e}")

        # Cost summary table
        self._add_subheader('Idle Time Cost Summary')
        summary = self.content['idle_summary_by_contractor']

        table = self.doc.add_table(rows=len(summary) + 2, cols=4)
        set_table_borders(table)

        headers = ['Contractor', 'Observations', 'Worker-Hours', f'Cost @ ${self.content["labor_rate"]}/hr']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            add_cell_shading(cell, 'D9D9D9')

        for idx, s in enumerate(summary, 1):
            table.rows[idx].cells[0].text = s['contractor']
            table.rows[idx].cells[1].text = str(s['observations'])
            table.rows[idx].cells[2].text = f"{s['worker_hours']:.0f}"
            table.rows[idx].cells[3].text = f"${s['cost']:,.0f}"

        # Total
        total_idx = len(summary) + 1
        table.rows[total_idx].cells[0].text = 'TOTAL'
        table.rows[total_idx].cells[1].text = str(total_obs)
        table.rows[total_idx].cells[2].text = f"{total_hours:.0f}"
        table.rows[total_idx].cells[3].text = f"${total_cost:,.0f}"
        for cell in table.rows[total_idx].cells:
            cell.paragraphs[0].runs[0].font.bold = True

    def _add_category_4(self):
        """Category 4: High Verification Locations."""
        self._add_category_header(4, 'High Verification Locations - Best Practices')

        narratives = self.content.get('narratives', {}).get('cat4', {})
        intro = narratives.get('intro', 'Several locations showed excellent verification rates, demonstrating proper TBM implementation.')

        p = self.doc.add_paragraph()
        p.add_run(intro)

        self.doc.add_paragraph()

        # Table
        high_locs = self.content['high_verification_locations']
        if high_locs:
            table = self.doc.add_table(rows=len(high_locs) + 1, cols=5)
            set_table_borders(table)

            headers = ['Contractor', 'Location', 'TBM', 'Verified', 'Accuracy']
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                add_cell_shading(cell, 'D9D9D9')

            for idx, loc in enumerate(high_locs, 1):
                table.rows[idx].cells[0].text = loc['company']
                table.rows[idx].cells[1].text = loc['location']
                table.rows[idx].cells[2].text = f"{loc['tbm_planned']:.0f}"
                table.rows[idx].cells[3].text = f"{loc['verified']:.0f}"
                table.rows[idx].cells[4].text = f"{loc['accuracy']:.0f}%"
                # Green font for high accuracy instead of background
                table.rows[idx].cells[4].paragraphs[0].runs[0].font.color.rgb = Colors.SUCCESS_GREEN
                table.rows[idx].cells[4].paragraphs[0].runs[0].font.bold = True

        self.doc.add_paragraph()

        note = narratives.get('note', 'Note: Accuracy >100% indicates more workers verified than committed in TBM.')
        p = self.doc.add_paragraph()
        run = p.add_run(note)
        run.font.italic = True
        run.font.size = Pt(11)

    def _add_category_5(self):
        """Category 5: Key Findings."""
        self._add_category_header(5, 'Key Findings')

        for finding in self.content['key_findings']:
            p = self.doc.add_paragraph(finding, style='List Bullet')

    def _add_category_6(self):
        """Category 6: Contractor Performance Comparison."""
        self._add_category_header(6, 'Contractor Performance Comparison')

        for m in self.content['contractor_metrics']:
            self._add_subheader(f"{m['name']} Performance")

            table = self.doc.add_table(rows=7, cols=2)
            set_table_borders(table)

            metrics = [
                ('Total Tasks', str(m['total_tasks'])),
                ('Total Planned', f"{m['total_planned']:.0f}"),
                ('Total Verified', f"{m['total_verified']:.0f}"),
                ('Accuracy', f"{m['accuracy']:.0f}%"),
                ('Idle Locations', str(m['idle_locations'])),
                ('Idle Cost', f"${m['idle_cost']:,.0f}"),
                ('Zero Verification', str(m['zero_verification_count'])),
            ]

            for idx, (label, value) in enumerate(metrics):
                table.rows[idx].cells[0].text = label
                table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
                add_cell_shading(table.rows[idx].cells[0], 'D9D9D9')
                table.rows[idx].cells[1].text = value

    def _add_category_7(self):
        """Category 7: TBM Process Limitations."""
        self._add_category_header(7, 'TBM Process Limitations')

        narratives = self.content.get('narratives', {}).get('cat7', {})

        # Default root cause based on actual data
        default_root_cause = (
            f"The low verification rates documented in this report ({self.content['lpi']:.0f}% overall, "
            f"{self.content['zero_verification_count']} locations with 0% verification) may be partially "
            f"attributable to verification time constraints."
        )

        self._add_subheader('7.1 Root Cause Analysis')
        root_cause = narratives.get('root_cause', default_root_cause)
        p = self.doc.add_paragraph()
        p.add_run(root_cause)

        self._add_subheader('7.2 Impact on Report Accuracy')
        default_caveats = [
            'Unverified locations may reflect time constraints, not actual absence',
            'Idle time observations limited to accessible areas',
            'Some high-verification locations may indicate concentrated verification',
            'Zero-verification locations may be due to inability to reach remote areas in time',
            'Actual contractor performance may be better than reported metrics suggest',
        ]
        caveats = narratives.get('caveats', default_caveats)
        for caveat in caveats:
            self.doc.add_paragraph(caveat, style='List Bullet')

        self._add_subheader('7.3 Required Process Improvements')
        default_improvements = [
            'TBM submission deadline: 8:00 AM daily (provides 6+ hour verification window)',
            'Minimum verification window: 6 hours from TBM receipt to worker departure',
            'Worker commitment: If in TBM, workers must remain on site until 4:00 PM minimum',
            'Real-time updates: Contractors must notify MXI immediately of schedule changes',
        ]
        improvements = narratives.get('improvements', default_improvements)
        for imp in improvements:
            self.doc.add_paragraph(imp, style='List Bullet')

    def _add_category_8(self):
        """Category 8: Conclusions & Recommendations."""
        self._add_category_header(8, 'Conclusions & Recommendations')

        narratives = self.content.get('narratives', {}).get('cat8', {})

        # Build default issues from actual data
        zero_rate = self.content['zero_verification_rate']
        idle_cost = self.content['idle_time_cost']
        default_issues = [
            f"TBM Documentation: {zero_rate:.0f}% of locations show zero verification",
            f"Idle time represents actual quantifiable waste (${idle_cost:,.0f}/day)",
        ]

        self._add_subheader('8.1 Primary Issues')
        issues = narratives.get('issues', default_issues)
        for issue in issues:
            self.doc.add_paragraph(issue, style='List Bullet')

        # Build default recommendations
        monthly_proj = self.content['monthly_projection']
        default_recommendations = [
            'CRITICAL: Require TBM submission by 8:00 AM to allow proper verification planning',
            'Establish minimum notice period: 6 hours between TBM submission and verification requirement',
            'Enhance TBM updates: Real-time location updates when workers reassigned',
            f'Continue monitoring idle time: ${monthly_proj:.0f}K monthly projection requires mitigation',
        ]

        self._add_subheader('8.2 Recommendations')
        recommendations = narratives.get('recommendations', default_recommendations)
        for rec in recommendations:
            p = self.doc.add_paragraph(style='List Bullet')
            if rec.startswith('CRITICAL:'):
                run = p.add_run(rec)
                run.font.bold = True
                run.font.color.rgb = Colors.CRITICAL_RED
            else:
                p.add_run(rec)

    def _add_end_marker(self):
        """Add end of report marker."""
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('— End of Report —')
        run.font.italic = True
        run.font.size = Pt(11)


# =============================================================================
# CLI
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description='Generate TBM Analysis DOCX')
    parser.add_argument('--input', required=True, help='Input JSON path')
    parser.add_argument('--output', help='Output DOCX path')

    args = parser.parse_args()

    # Load content
    input_path = Path(args.input)
    with open(input_path) as f:
        content = json.load(f)

    # Determine output path
    if args.output:
        output_path = args.output
    else:
        date_str = content.get('report_date', 'unknown')
        group = content.get('contractor_group', 'TBM')
        # Convert date format
        if len(date_str) == 8:
            parts = date_str.split('-')
            date_fmt = f"{parts[0]}.{parts[1]}.{parts[2]}"
        else:
            try:
                dt = datetime.strptime(date_str, '%Y-%m-%d')
                date_fmt = dt.strftime('%m.%d.%y')
            except:
                date_fmt = date_str

        output_path = input_path.parent / f"[GENERATED] MXI - {group} TBM Analysis - {date_fmt}.docx"

    # Build report
    print(f"Generating: {output_path}")
    builder = TBMReportBuilder(content)
    builder.build(str(output_path))
    print(f"Done: {output_path}")


if __name__ == '__main__':
    main()
